from docx import Document

def docx_to_txt(docx_path, txt_path):
    # 加载Word文档
    doc = Document(docx_path)
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        # 遍历文档中的每一段
        for paragraph in doc.paragraphs:
            txt_file.write(paragraph.text + '\n')  # 写入段落文本

        # 遍历文档中的每个表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt_file.write(cell.text + '\t')  # 使用制表符分隔单元格内容
                txt_file.write('\n')  # 每行结束后换行
            txt_file.write('\n')  # 每个表格结束后再空一行

    return

def docx_to_markdown(docx_path, md_path):
    doc = Document(docx_path)
    with open(md_path, 'w', encoding='utf-8') as md_file:
        # 遍历文档中的每一段
        for paragraph in doc.paragraphs:
            text = paragraph.text
            md_file.write(text + '\n\n')

        # 遍历文档中的每个表格
        for table in doc.tables:
            for row_index, row in enumerate(table.rows):
                row_text = '|'
                for cell in row.cells:
                    row_text += f" {cell.text} |"
                md_file.write(row_text + '\n')

                # 如果是第一行，则添加分隔行
                if row_index == 0:
                    md_file.write('|' + ' --- |' * len(row.cells) + '\n')
            md_file.write('\n')  # 每个表格结束后空一行
    return

# 调用函数，传入Word文档路径和想要保存的txt文件路径
docx_to_txt('..\\..\\data\\workflows\\Contracts\\input\\raw_data\\空调合同（样例）-训练用.docx',
            '..\\..\\data\\workflows\\Contracts\\input\\text_data\\contract_5.txt')
